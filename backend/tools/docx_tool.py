from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_docx(content: str, style_params: dict, file_path: str) -> str:
    """Generate a Word document from content with style parameters."""

    doc = Document()

    font_family = style_params.get("font_family", "Calibri")
    font_size = style_params.get("font_size", 12)
    text_color = style_params.get("text_color", "000000")
    line_spacing = style_params.get("line_spacing", 1.0)

    paragraphs = content.split("\n\n")

    for i, para_text in enumerate(paragraphs):
        if not para_text.strip():
            continue

        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.line_spacing = line_spacing
        paragraph.paragraph_format.space_after = Pt(12)

        if i == 0:
            run = paragraph.add_run(para_text.strip())
            run.font.size = Pt(font_size + 4)
            run.font.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            run = paragraph.add_run(para_text.strip())
            run.font.size = Pt(font_size)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        run.font.name = font_family
        run.font.color.rgb = RGBColor(*_hex_to_rgb(text_color))

    doc.save(file_path)
    return file_path


def _hex_to_rgb(hex_color: str):
    """Convert hex color to RGB tuple."""
    hex_color = hex_color.lstrip("#")
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
